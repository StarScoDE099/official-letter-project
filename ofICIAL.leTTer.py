from docx import Document

document = Document()
# SENDER ADDRESS
p = document.add_paragraph()
First_name = input('What is your First name? ')
Surname = input('What is your Surname? ')
Full_name = First_name + ' ' + Surname
address = input('Enter Address ')
location = input('Enter location ')


p.add_run(Full_name + '\n ' + address + '\n ' + location ).bold = True

p = document.add_paragraph()
Date = input('Enter Date ')

p.add_run(Date + ' ').bold = True

# INSTITUTION ADDRESS.
document.add_heading('Through')
p = document.add_paragraph()

Recognition = input('Enter REcognition ')
Institution = input('Enter Institution ')
address2 = input('Enter Address ')
location2 = input('Enter location ')

p.add_run(Recognition + '\n' + Institution + '\n' + address2 + '\n' + location2 ).bold =True

# COMPANY ADDRESS.
document.add_heading('To ,')
p = document.add_paragraph()

Reciver =input('Enter Reciver ')

p.add_run(Reciver + ' ').bold = True


document.add_heading('Dear Sir / Madam. ')
p=document.add_paragraph ()

# REMAX
document.add_heading('REF : REQUEST FOR CYBER SECURITY ATTACHMENT. ')
p = document.add_paragraph()

Txt =input('Enter Txt. ')

p=document.add_paragraph(Txt)

#End REMAX

Remax = input('Enter Remax ')
conclution = input ('Enter conclution ')
Sighn = input('Enter Sighn')
name3 = input('Enter name3 ')
Phone_Number =input('What is your phone number?' )

p = document.add_paragraph(
    Remax + '\n' + conclution + '\n' + name3 + '\n' + Sighn + '\n' + Phone_Number )






















document.save('oFficial.docx')